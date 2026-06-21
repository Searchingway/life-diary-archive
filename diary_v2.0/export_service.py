from __future__ import annotations

import json
import re
import sys
import zipfile
from datetime import date, datetime
from pathlib import Path
from typing import Any

_src = str((Path(__file__).resolve().parent.parent / "src").resolve())
if _src not in sys.path:
    sys.path.insert(0, _src)

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt

from life_dairy.exporters import DiaryExporter, DiaryExportItem
from life_dairy.models import DiaryEntry, DiaryImage

from data_api import (
    MODULE_BY_KEY,
    MODULES,
    build_diary_export_item,
    configured_export_dir,
    DATA_ROOT,
    footprint_image_file,
    image_metadata,
    list_module_records,
    module_record_export_lines,
    now_iso,
    safe_export_name,
    unique_output_dir,
    unique_output_path,
)


def set_docx_runs_font(paragraph: Any, size: int = 11, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(size)
        run.bold = bold


def add_docx_paragraph(document: Any, text: str = "", size: int = 11, bold: bool = False) -> Any:
    paragraph = document.add_paragraph(text)
    set_docx_runs_font(paragraph, size=size, bold=bold)
    return paragraph


def append_footprint_record_to_docx(document: Any, record: dict[str, Any], include_page_break: bool = False) -> None:
    if include_page_break:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    place_id = str(record.get("id") or "")
    title = str(record.get("title") or "未命名足迹")
    heading = document.add_heading(title, level=1)
    set_docx_runs_font(heading, size=18, bold=True)
    body = str(record.get("body") or "").strip()
    visits = record.get("extra", {}).get("visits", []) if isinstance(record.get("extra"), dict) else []
    add_docx_paragraph(document, f"地点记录数：{len(visits)}", size=10)
    if record.get("date"):
        add_docx_paragraph(document, f"日期：{record.get('date')}", size=10)
    if body:
        add_docx_paragraph(document, "地点描述", size=13, bold=True)
        for line in body.splitlines():
            add_docx_paragraph(document, line, size=11)
    if not visits:
        add_docx_paragraph(document, "（暂无访问记录）", size=11)
        return
    for visit in visits:
        if not isinstance(visit, dict):
            continue
        visit_id = str(visit.get("id") or "")
        visit_date = str(visit.get("date") or visit.get("visit_date") or "未设置日期")
        date_heading = document.add_heading(visit_date, level=2)
        set_docx_runs_font(date_heading, size=14, bold=True)
        thought = str(visit.get("thought") or visit.get("reflection") or "").strip()
        if thought:
            for line in thought.splitlines():
                add_docx_paragraph(document, line, size=11)
        else:
            add_docx_paragraph(document, "（本次访问暂无文字记录）", size=11)
        images = visit.get("images") if isinstance(visit.get("images"), list) else []
        exportable_images: list[tuple[str, Path]] = []
        for image in images:
            parsed = image_metadata(image)
            if not parsed:
                continue
            file_name, label = parsed
            image_path = footprint_image_file(place_id, visit_id, file_name)
            if image_path and image_path.exists():
                exportable_images.append((label or file_name, image_path))
        if exportable_images:
            add_docx_paragraph(document, "图片", size=12, bold=True)
            for caption_text, image_path in exportable_images:
                table = document.add_table(rows=1, cols=1)
                cell = table.cell(0, 0)
                caption = cell.paragraphs[0]
                caption.add_run(caption_text)
                set_docx_runs_font(caption, size=10, bold=True)
                image_paragraph = cell.add_paragraph()
                try:
                    image_paragraph.add_run().add_picture(str(image_path), width=Cm(14.8))
                except Exception:
                    image_paragraph.add_run("（图片无法加载）")
                document.add_paragraph("")


def zip_raw_data(target_dir: Path) -> Path:
    zip_path = unique_output_path(target_dir / "2.0原始数据.zip")
    exports_root = (DATA_ROOT / "exports").resolve()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in DATA_ROOT.rglob("*"):
            if not path.is_file():
                continue
            resolved = path.resolve()
            if str(resolved).startswith(str(exports_root)):
                continue
            archive.write(path, path.relative_to(DATA_ROOT))
    return zip_path


def export_entry_word_pdf(entry_id: str) -> dict[str, str]:
    record_dir = DATA_ROOT / "entries" / entry_id
    metadata_path = record_dir / "entry.json"
    if not metadata_path.exists():
        raise FileNotFoundError(f"entry not found: {entry_id}")
    data = __import__("json").loads(metadata_path.read_text(encoding="utf-8"))
    body = (record_dir / str(data.get("body_file", "content.md"))).read_text(encoding="utf-8", errors="replace")
    images = []
    image_paths: list[Path] = []
    for value in data.get("images", []):
        parsed = image_metadata(value)
        if not parsed:
            continue
        file_name, label = parsed
        images.append(DiaryImage(file_name=file_name, label=label))
        path = record_dir / "images" / file_name
        if path.exists():
            image_paths.append(path)
    entry = DiaryEntry(
        id=str(data.get("id") or entry_id),
        date=str(data.get("date") or date.today().isoformat()),
        title=str(data.get("title") or ""),
        body=body,
        created_at=str(data.get("created_at") or now_iso()),
        updated_at=str(data.get("updated_at") or now_iso()),
        images=images,
    )
    export_dir = configured_export_dir()
    exporter = DiaryExporter(export_dir)
    docx_path, pdf_path = exporter.export_entries_word_and_pdf(
        [DiaryExportItem(entry=entry, image_lookup=exporter._build_image_lookup(image_paths))]
    )
    return {"docx_path": str(docx_path), "pdf_path": str(pdf_path)}


def export_all_entries_word_pdf(payload: dict[str, Any] | None = None) -> dict[str, Any]:
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)
    export_items: list[DiaryExportItem] = []
    entries_dir = DATA_ROOT / "entries"
    entries_dir.mkdir(parents=True, exist_ok=True)
    for child in entries_dir.iterdir():
        if not child.is_dir():
            continue
        item = build_diary_export_item(child)
        if item is not None:
            export_items.append(item)
    if not export_items:
        raise ValueError("没有可导出的日记")
    exporter = DiaryExporter(output_dir)
    docx_path, pdf_path = exporter.export_entries_word_and_pdf(export_items, export_all=True)
    return {"docx_path": str(docx_path), "pdf_path": str(pdf_path), "output_dir": str(output_dir), "count": len(export_items)}


def export_all_entries_txt(payload: dict[str, Any] | None = None) -> dict[str, Any]:
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)
    export_items: list[DiaryExportItem] = []
    entries_dir = DATA_ROOT / "entries"
    entries_dir.mkdir(parents=True, exist_ok=True)
    for child in entries_dir.iterdir():
        if not child.is_dir():
            continue
        item = build_diary_export_item(child)
        if item is not None:
            export_items.append(item)
    if not export_items:
        raise ValueError("没有可导出的日记")
    export_items.sort(key=lambda item: (item.entry.date, item.entry.created_at, item.entry.updated_at), reverse=True)
    txt_path = unique_output_path(output_dir / f"{safe_export_name(f'全部日记_{len(export_items)}篇日记')}.txt")
    lines: list[str] = []
    for index, item in enumerate(export_items, start=1):
        entry = item.entry
        body = entry.body.strip()
        lines.extend([
            f"{index}. {entry.display_title}",
            f"日期：{entry.date}",
            f"字数：{len(body)}",
            "",
            body or "（正文为空）",
            "",
            "-" * 48,
            "",
        ])
    txt_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return {"txt_path": str(txt_path), "output_dir": str(output_dir), "count": len(export_items)}


def export_module_txt(module_key: str, payload: dict[str, Any] | None = None) -> dict[str, Any]:
    if module_key not in MODULE_BY_KEY:
        raise ValueError(f"未知导出模块：{module_key}")
    module = MODULE_BY_KEY[module_key]
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)
    records = list_module_records(module_key)
    if not records:
        raise ValueError(f"没有可导出的{module.label}记录")
    txt_path = unique_output_path(output_dir / f"{safe_export_name(f'{module.label}_{len(records)}条记录')}.txt")
    lines = [f"{module.label}导出", f"导出时间：{now_iso()}", f"记录数：{len(records)}", ""]
    for index, record in enumerate(records, start=1):
        lines.extend(module_record_export_lines(module, record, index))
    txt_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return {"module_key": module.key, "module_label": module.label, "txt_path": str(txt_path), "output_dir": str(output_dir), "count": len(records)}


def export_notes_markdown(payload: dict[str, Any] | None = None) -> dict[str, Any]:
    module = MODULE_BY_KEY["notes"]
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    target_dir = unique_output_dir(output_dir / f"{safe_export_name(f'{module.label}_Markdown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    target_dir.mkdir(parents=True, exist_ok=True)
    records = list_module_records("notes")
    if not records:
        raise ValueError("没有可导出的信息备忘记录")
    files: list[dict[str, Any]] = []
    for index, record in enumerate(records, start=1):
        title = str(record.get("title") or f"信息备忘_{index}").strip()
        md_path = unique_output_path(target_dir / f"{index:03d}_{safe_export_name(title)}.md")
        front_matter = ["---", f"title: {title}", f"date: {record.get('date') or ''}", f"updated_at: {record.get('updated_at') or ''}", "---", ""]
        body = str(record.get("body") or "").rstrip()
        md_path.write_text("\n".join(front_matter) + body + "\n", encoding="utf-8")
        files.append({"module_key": module.key, "module_label": module.label, "md_path": str(md_path), "count": 1})
    return {"module_key": module.key, "module_label": module.label, "output_dir": str(target_dir), "count": len(records), "files": files}


def export_footprints_word(payload: dict[str, Any] | None = None, output_dir: Path | None = None) -> dict[str, Any]:
    target_dir = output_dir
    if target_dir is None:
        output_text = str((payload or {}).get("output_dir") or "").strip()
        target_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    target_dir.mkdir(parents=True, exist_ok=True)
    records = list_module_records("footprints")
    if not records:
        raise ValueError("没有可导出的足迹记录")
    docx_path = unique_output_path(target_dir / f"{safe_export_name(f'足迹_{len(records)}条记录')}.docx")
    document = Document()
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"].font.size = Pt(11)
    title = document.add_heading("足迹导出", level=0)
    set_docx_runs_font(title, size=20, bold=True)
    add_docx_paragraph(document, f"导出时间：{now_iso()}", size=10)
    add_docx_paragraph(document, f"地点数：{len(records)}", size=10)
    for index, record in enumerate(records, start=1):
        append_footprint_record_to_docx(document, record, include_page_break=index > 1)
    document.save(docx_path)
    return {"module_key": "footprints", "module_label": MODULE_BY_KEY["footprints"].label, "docx_path": str(docx_path), "output_dir": str(target_dir), "count": len(records)}


def export_all_modules(payload: dict[str, Any] | None = None) -> dict[str, Any]:
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)
    package_dir = unique_output_dir(output_dir / f"全部板块导出_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    package_dir.mkdir(parents=True, exist_ok=True)
    files: list[dict[str, Any]] = []
    total = 0
    for module in MODULES:
        records = list_module_records(module.key)
        txt_path = package_dir / f"{safe_export_name(module.label)}.txt"
        lines = [f"{module.label}导出", f"导出时间：{now_iso()}", f"记录数：{len(records)}", ""]
        if records:
            for index, record in enumerate(records, start=1):
                lines.extend(module_record_export_lines(module, record, index))
        else:
            lines.append("暂无数据")
        txt_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        total += len(records)
        files.append({"module_key": module.key, "module_label": module.label, "txt_path": str(txt_path), "count": len(records)})
        if module.key == "footprints" and records:
            footprint_word = export_footprints_word(output_dir=package_dir)
            files.append({"module_key": module.key, "module_label": module.label, "docx_path": footprint_word["docx_path"], "count": len(records)})
    zip_path = zip_raw_data(package_dir)
    manifest_path = package_dir / "导出说明.json"
    manifest_path.write_text(
        json.dumps(
            {"exported_at": now_iso(), "data_root": str(DATA_ROOT), "total_records": total, "files": files, "raw_data_zip": str(zip_path)},
            ensure_ascii=False, indent=2,
        ), encoding="utf-8",
    )
    return {"output_dir": str(package_dir), "count": total, "files": files, "zip_path": str(zip_path), "manifest_path": str(manifest_path)}

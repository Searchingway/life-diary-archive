from __future__ import annotations

import json
import mimetypes
import os
import base64
import re
import shutil
import subprocess
import sys
import threading
import uuid
import webbrowser
import zipfile
from dataclasses import dataclass
from datetime import date, datetime
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, unquote, urlparse

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt

sys.path.insert(0, str((Path(__file__).resolve().parent.parent / "src").resolve()))

from life_dairy.exporters import DiaryExporter, DiaryExportItem
from life_dairy.models import DiaryEntry, DiaryImage

from PySide6.QtCore import QUrl
from PySide6.QtWidgets import QApplication, QMainWindow
from PySide6.QtWebEngineWidgets import QWebEngineView


APP_TITLE = "人生档案 Diary Desktop 2.0"
BASE_DIR = Path(__file__).resolve().parent
REPO_ROOT = BASE_DIR.parent
FRONTEND_DIST = BASE_DIR / "Untitled" / "dist"
LEGACY_DATA_ROOT = (REPO_ROOT / "data" / "Diary").resolve()
DATA_ROOT = Path(os.environ.get("LIFE_DIARY_DATA_ROOT", BASE_DIR / "data" / "Diary")).resolve()
MIGRATION_MARKER = DATA_ROOT / ".life_diary_2_migration.json"
SETTINGS_PATH = DATA_ROOT / "config" / "app_settings.json"


@dataclass(frozen=True)
class ModuleConfig:
    key: str
    label: str
    directory: str
    json_file: str
    title_fields: tuple[str, ...]
    body_files: tuple[str, ...] = ()
    body_fields: tuple[str, ...] = ()
    date_fields: tuple[str, ...] = ("date", "due_date", "start_date", "created_at")
    type_fields: tuple[str, ...] = ("type", "plan_type", "work_type", "analysis_type", "info_type")
    status_fields: tuple[str, ...] = ("status", "priority")


MODULES: list[ModuleConfig] = [
    ModuleConfig("entries", "日记", "entries", "entry.json", ("title",), ("content.md",)),
    ModuleConfig("footprints", "足迹", "footprints", "footprint.json", ("place_name",), ("summary.md",)),
    ModuleConfig("plans", "轻计划", "plans", "plan.json", ("title",), body_fields=("notes", "reason", "alternative_action")),
    ModuleConfig("action_plans", "行动计划", "action_plans", "action_plan.json", ("title",), body_fields=("description", "summary")),
    ModuleConfig("thoughts", "轻思考", "thoughts", "thought.json", ("title",), body_fields=("description", "preliminary_conclusion", "notes")),
    ModuleConfig("resources", "轻资源", "resources", "resource.json", ("title",), body_fields=("description", "overall_judgement", "notes")),
    ModuleConfig("info_memos", "信息备忘", "info_memos", "info_memo.json", ("title",), body_fields=("main_content", "notes", "description")),
    ModuleConfig("observations", "自我观察", "observations", "observation.json", ("title", "scene"), body_fields=("description", "body", "notes")),
    ModuleConfig("lessons", "教训与反思", "lessons", "lesson.json", ("title",), ("event.md", "reflection.md"), ("summary", "cost")),
    ModuleConfig("self_analysis", "自我分析", "self_analysis", "analysis.json", ("title",), ("content.md",), ("summary", "analysis")),
    ModuleConfig("works", "作品感悟", "works", "work.json", ("title", "work_title"), ("content.md", "summary.md"), ("summary", "notes")),
    ModuleConfig("notes", "笔记", "notes", "note.json", ("title",), ("content.md",), ("content", "notes")),
]

MODULE_BY_KEY = {module.key: module for module in MODULES}


def now_iso() -> str:
    return datetime.now().astimezone().isoformat(timespec="microseconds")


def migrate_legacy_data_if_needed() -> bool:
    if os.environ.get("LIFE_DIARY_SKIP_MIGRATION") == "1":
        DATA_ROOT.mkdir(parents=True, exist_ok=True)
        return False
    if MIGRATION_MARKER.exists():
        return False
    if DATA_ROOT.exists() and any(DATA_ROOT.iterdir()):
        return False
    DATA_ROOT.mkdir(parents=True, exist_ok=True)
    if not LEGACY_DATA_ROOT.exists():
        MIGRATION_MARKER.write_text(
            json.dumps(
                {
                    "migrated": False,
                    "reason": "legacy data root not found",
                    "legacy_data_root": str(LEGACY_DATA_ROOT),
                    "created_at": now_iso(),
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        return False
    for child in LEGACY_DATA_ROOT.iterdir():
        target = DATA_ROOT / child.name
        if child.is_dir():
            shutil.copytree(child, target, dirs_exist_ok=True)
        elif child.is_file():
            shutil.copy2(child, target)
    MIGRATION_MARKER.write_text(
        json.dumps(
            {
                "migrated": True,
                "legacy_data_root": str(LEGACY_DATA_ROOT),
                "data_root": str(DATA_ROOT),
                "created_at": now_iso(),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return True


def first_text(data: dict[str, Any], fields: tuple[str, ...]) -> str:
    for field in fields:
        value = data.get(field)
        if value is None:
            continue
        text = str(value).strip()
        if text:
            return text
    return ""


def read_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as file:
        value = json.load(file)
    return value if isinstance(value, dict) else {}


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def read_settings() -> dict[str, Any]:
    if SETTINGS_PATH.exists():
        try:
            return read_json(SETTINGS_PATH)
        except (OSError, json.JSONDecodeError):
            return {}
    return {}


def write_settings(data: dict[str, Any]) -> dict[str, Any]:
    settings = {**read_settings(), **data, "updated_at": now_iso()}
    write_json(SETTINGS_PATH, settings)
    return settings


def default_export_dir() -> Path:
    return (DATA_ROOT / "exports").resolve()


def configured_export_dir() -> Path:
    value = str(read_settings().get("export_dir") or "").strip()
    output_dir = Path(value).expanduser().resolve() if value else default_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def read_body(record_dir: Path, data: dict[str, Any], module: ModuleConfig) -> str:
    body_parts: list[str] = []
    candidates = list(module.body_files)
    body_file = data.get("body_file")
    if isinstance(body_file, str) and body_file not in candidates:
        candidates.insert(0, body_file)

    for filename in candidates:
        path = record_dir / filename
        if path.exists() and path.is_file():
            text = path.read_text(encoding="utf-8", errors="replace").strip()
            if text:
                body_parts.append(text)

    for field in module.body_fields:
        value = data.get(field)
        if value is None:
            continue
        text = str(value).strip()
        if text:
            body_parts.append(text)

    if module.key == "resources":
        for item in data.get("resource_items", []):
            if isinstance(item, dict):
                item_type = str(item.get("type") or item.get("label") or "").strip()
                item_value = str(item.get("value") or item.get("description") or "").strip()
                if item_type or item_value:
                    body_parts.append(f"{item_type}: {item_value}".strip(": "))

    return "\n\n".join(dict.fromkeys(body_parts))


def write_body(record_dir: Path, data: dict[str, Any], module: ModuleConfig, body: str) -> None:
    if module.body_files:
        filename = module.body_files[0]
        if module.key == "entries":
            data["body_file"] = filename
        elif module.key in {"footprints", "self_analysis", "works", "notes"}:
            if filename.endswith(".md"):
                field = "summary_file" if filename == "summary.md" else "content_file"
                data.setdefault(field, filename)
        (record_dir / filename).write_text(body, encoding="utf-8")
        return
    if module.body_fields:
        data[module.body_fields[0]] = body


def record_from_directory(module: ModuleConfig, record_dir: Path) -> dict[str, Any] | None:
    json_path = record_dir / module.json_file
    if not json_path.exists():
        return None
    try:
        data = read_json(json_path)
    except (OSError, json.JSONDecodeError):
        return None
    if data.get("deleted"):
        return None

    record_id = str(data.get("id") or record_dir.name)
    title = first_text(data, module.title_fields) or "未命名记录"
    date_text = first_text(data, module.date_fields)
    updated_at = str(data.get("updated_at") or data.get("created_at") or "")
    record_type = first_text(data, module.type_fields)
    status = first_text(data, module.status_fields)
    subtitle = " / ".join(part for part in (record_type, status) if part) or module.label
    if module.key == "footprints":
        data = {**data, "visits": read_footprint_visits(record_dir, data)}

    return {
        "id": record_id,
        "title": title,
        "subtitle": subtitle,
        "body": read_body(record_dir, data, module),
        "date": date_text[:10] if len(date_text) >= 10 else date_text,
        "updated_at": updated_at,
        "status": status,
        "type": record_type,
        "extra": data,
    }


def list_module_records(module_key: str, query: str = "") -> list[dict[str, Any]]:
    module = MODULE_BY_KEY[module_key]
    module_dir = DATA_ROOT / module.directory
    module_dir.mkdir(parents=True, exist_ok=True)
    keyword = query.strip().lower()
    records: list[dict[str, Any]] = []

    for child in module_dir.iterdir():
        if not child.is_dir():
            continue
        record = record_from_directory(module, child)
        if not record:
            continue
        haystack = "\n".join(str(record.get(key, "")) for key in ("title", "subtitle", "body", "date")).lower()
        if keyword and keyword not in haystack:
            continue
        records.append(record)

    if module_key == "entries":
        records.sort(
            key=lambda item: (str(item.get("date") or ""), str(item.get("updated_at") or "")),
            reverse=True,
        )
    else:
        records.sort(key=lambda item: str(item.get("updated_at") or item.get("date") or ""), reverse=True)
    return records


def build_overview() -> dict[str, Any]:
    modules: list[dict[str, Any]] = []
    recent: list[dict[str, Any]] = []
    for module in MODULES:
        records = list_module_records(module.key)
        latest = records[0]["updated_at"][:10] if records and records[0].get("updated_at") else ""
        modules.append({"key": module.key, "label": module.label, "count": len(records), "latest": latest})
        for record in records[:5]:
            recent.append({**record, "module": module.label, "module_key": module.key})

    recent.sort(key=lambda item: str(item.get("updated_at") or item.get("date") or ""), reverse=True)
    return {
        "data_root": str(DATA_ROOT),
        "legacy_data_root": str(LEGACY_DATA_ROOT),
        "migrated_from_legacy": MIGRATION_MARKER.exists(),
        "dashboard_stats": build_dashboard_stats(),
        "modules": modules,
        "recent": recent[:20],
    }


def build_dashboard_stats() -> dict[str, int]:
    today = date.today()
    month_prefix = today.strftime("%Y-%m")
    year_prefix = today.strftime("%Y")
    entries = list_module_records("entries")
    light_plans = list_module_records("plans")
    action_plans = list_module_records("action_plans")

    month_entries = [entry for entry in entries if str(entry.get("date", "")).startswith(month_prefix)]
    year_entries = [entry for entry in entries if str(entry.get("date", "")).startswith(year_prefix)]

    def word_count(records: list[dict[str, Any]]) -> int:
        return sum(len(str(record.get("body", ""))) for record in records)

    def image_count(records: list[dict[str, Any]]) -> int:
        total = 0
        for record in records:
            extra = record.get("extra")
            if isinstance(extra, dict) and isinstance(extra.get("images"), list):
                total += len(extra["images"])
        return total

    def completed(record: dict[str, Any]) -> bool:
        status = str(record.get("status") or "")
        extra = record.get("extra") if isinstance(record.get("extra"), dict) else {}
        tasks = extra.get("tasks") if isinstance(extra, dict) else []
        if "完成" in status or status.lower() in {"completed", "done"}:
            return True
        if isinstance(tasks, list) and tasks:
            return all(isinstance(task, dict) and task.get("done") for task in tasks)
        return False

    def active(record: dict[str, Any]) -> bool:
        status = str(record.get("status") or "")
        return ("进行" in status) or status.lower() in {"active", "in_progress"} or not completed(record)

    def updated_in(record: dict[str, Any], prefix: str) -> bool:
        value = str(record.get("updated_at") or record.get("date") or "")
        return value.startswith(prefix)

    today_str = today.isoformat()
    today_pending = 0
    for plan in action_plans:
        extra = plan.get("extra") if isinstance(plan.get("extra"), dict) else {}
        tasks = extra.get("tasks") if isinstance(extra, dict) else []
        if isinstance(tasks, list):
            today_pending += sum(
                1
                for task in tasks
                if isinstance(task, dict)
                and str(task.get("date") or "") == today_str
                and not bool(task.get("done"))
            )

    return {
        "month_diary_count": len(month_entries),
        "month_diary_words": word_count(month_entries),
        "month_diary_images": image_count(month_entries),
        "month_completed_plans": sum(1 for plan in light_plans + action_plans if completed(plan) and updated_in(plan, month_prefix)),
        "year_diary_count": len(year_entries),
        "year_diary_words": word_count(year_entries),
        "year_diary_images": image_count(year_entries),
        "year_completed_plans": sum(1 for plan in light_plans + action_plans if completed(plan) and updated_in(plan, year_prefix)),
        "action_plan_count": len(action_plans),
        "active_action_plan_count": sum(1 for plan in action_plans if active(plan)),
        "today_pending_tasks": today_pending,
    }


def save_entry(payload: dict[str, Any]) -> dict[str, Any]:
    record_id = str(payload.get("id") or uuid.uuid4().hex)
    record_dir = DATA_ROOT / "entries" / record_id
    record_dir.mkdir(parents=True, exist_ok=True)
    metadata_path = record_dir / "entry.json"
    existing = read_json(metadata_path) if metadata_path.exists() else {}
    timestamp = now_iso()
    data = {
        **existing,
        "id": record_id,
        "date": str(payload.get("date") or date.today().isoformat()),
        "title": str(payload.get("title") or ""),
        "images": existing.get("images", []),
        "created_at": existing.get("created_at") or timestamp,
        "updated_at": timestamp,
        "body_file": "content.md",
    }
    metadata_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    (record_dir / "content.md").write_text(str(payload.get("body") or ""), encoding="utf-8")
    return record_from_directory(MODULE_BY_KEY["entries"], record_dir) or {}


def image_metadata(value: Any) -> tuple[str, str] | None:
    if isinstance(value, str):
        return value, ""
    if isinstance(value, dict):
        file_name = str(value.get("file_name") or value.get("name") or "")
        label = str(value.get("label") or "")
        if file_name:
            return file_name, label
    return None


def unique_image_name(images_dir: Path, original_name: str) -> str:
    candidate = Path(original_name).name or "image"
    stem = Path(candidate).stem or "image"
    suffix = Path(candidate).suffix or ".png"
    candidate = f"{stem}{suffix}"
    counter = 1
    while (images_dir / candidate).exists():
        candidate = f"{stem}_{counter}{suffix}"
        counter += 1
    return candidate


def add_entry_images(entry_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    record_dir = DATA_ROOT / "entries" / entry_id
    metadata_path = record_dir / "entry.json"
    if not metadata_path.exists():
        raise FileNotFoundError(f"entry not found: {entry_id}")
    images_dir = record_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    data = read_json(metadata_path)
    images = list(data.get("images", [])) if isinstance(data.get("images"), list) else []
    files = payload.get("files")
    if not isinstance(files, list):
        raise ValueError("files must be a list")

    for item in files:
        if not isinstance(item, dict):
            continue
        raw_data = str(item.get("data") or "")
        if "," in raw_data and raw_data.startswith("data:"):
            raw_data = raw_data.split(",", 1)[1]
        content = base64.b64decode(raw_data)
        file_name = unique_image_name(images_dir, str(item.get("name") or "image.png"))
        (images_dir / file_name).write_bytes(content)
        images.append({"file_name": file_name, "label": str(item.get("label") or "")})

    data["images"] = images
    data["updated_at"] = now_iso()
    write_json(metadata_path, data)
    return record_from_directory(MODULE_BY_KEY["entries"], record_dir) or {}


def entry_image_path(entry_id: str, image_name: str) -> Path:
    candidate = (DATA_ROOT / "entries" / entry_id / "images" / image_name).resolve()
    images_root = (DATA_ROOT / "entries" / entry_id / "images").resolve()
    if not str(candidate).startswith(str(images_root)):
        raise ValueError("invalid image path")
    if not candidate.exists():
        raise FileNotFoundError(f"image not found: {image_name}")
    return candidate


def export_entry_word_pdf(entry_id: str) -> dict[str, str]:
    record_dir = DATA_ROOT / "entries" / entry_id
    metadata_path = record_dir / "entry.json"
    if not metadata_path.exists():
        raise FileNotFoundError(f"entry not found: {entry_id}")
    data = read_json(metadata_path)
    body = (record_dir / str(data.get("body_file", "content.md"))).read_text(encoding="utf-8", errors="replace")
    images = []
    image_paths: list[Path] = []
    for value in data.get("images", []):
        parsed = image_metadata(value)
        if not parsed:
            continue
        file_name, label = parsed
        images.append(DiaryImage(file_name=file_name, label=label))
        path = record_dir / "images" / file_name
        if path.exists():
            image_paths.append(path)
    entry = DiaryEntry(
        id=str(data.get("id") or entry_id),
        date=str(data.get("date") or date.today().isoformat()),
        title=str(data.get("title") or ""),
        body=body,
        created_at=str(data.get("created_at") or now_iso()),
        updated_at=str(data.get("updated_at") or now_iso()),
        images=images,
    )
    export_dir = configured_export_dir()
    exporter = DiaryExporter(export_dir)
    docx_path, pdf_path = exporter.export_entries_word_and_pdf(
        [DiaryExportItem(entry=entry, image_lookup=exporter._build_image_lookup(image_paths))]
    )
    return {"docx_path": str(docx_path), "pdf_path": str(pdf_path)}


def build_diary_export_item(record_dir: Path) -> DiaryExportItem | None:
    metadata_path = record_dir / "entry.json"
    if not metadata_path.exists():
        return None
    data = read_json(metadata_path)
    if data.get("deleted"):
        return None
    body_path = record_dir / str(data.get("body_file", "content.md"))
    body = body_path.read_text(encoding="utf-8", errors="replace") if body_path.exists() else ""
    images: list[DiaryImage] = []
    image_paths: list[Path] = []
    for value in data.get("images", []):
        parsed = image_metadata(value)
        if not parsed:
            continue
        file_name, label = parsed
        images.append(DiaryImage(file_name=file_name, label=label))
        path = record_dir / "images" / file_name
        if path.exists():
            image_paths.append(path)
    entry = DiaryEntry(
        id=str(data.get("id") or record_dir.name),
        date=str(data.get("date") or date.today().isoformat()),
        title=str(data.get("title") or ""),
        body=body,
        created_at=str(data.get("created_at") or now_iso()),
        updated_at=str(data.get("updated_at") or now_iso()),
        images=images,
    )
    return DiaryExportItem(
        entry=entry,
        image_lookup={path.name: path for path in image_paths},
    )


def select_export_directory() -> Path:
    default_dir = configured_export_dir()
    default_dir.mkdir(parents=True, exist_ok=True)
    script = r"""
Add-Type -AssemblyName System.Windows.Forms
[Console]::OutputEncoding = [System.Text.Encoding]::UTF8
$dialog = New-Object System.Windows.Forms.FolderBrowserDialog
$dialog.Description = '选择日记导出目录'
$dialog.SelectedPath = $env:LIFE_DIARY_EXPORT_DEFAULT
$dialog.ShowNewFolderButton = $true
if ($dialog.ShowDialog() -eq [System.Windows.Forms.DialogResult]::OK) {
    Write-Output $dialog.SelectedPath
}
"""
    env = os.environ.copy()
    env["LIFE_DIARY_EXPORT_DEFAULT"] = str(default_dir)
    result = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        check=False,
        capture_output=True,
        text=True,
        timeout=300,
        env=env,
        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
    )
    selected = (result.stdout or "").strip().splitlines()
    if result.returncode != 0:
        raise RuntimeError((result.stderr or result.stdout or "选择导出目录失败").strip())
    if not selected:
        raise ValueError("已取消选择导出目录")
    output_dir = Path(selected[-1]).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def select_and_store_export_directory() -> dict[str, str]:
    output_dir = select_export_directory()
    write_settings({"export_dir": str(output_dir)})
    return {"export_dir": str(output_dir)}


def export_all_entries_word_pdf(payload: dict[str, Any] | None = None) -> dict[str, Any]:
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)

    export_items: list[DiaryExportItem] = []
    entries_dir = DATA_ROOT / "entries"
    entries_dir.mkdir(parents=True, exist_ok=True)
    for child in entries_dir.iterdir():
        if not child.is_dir():
            continue
        item = build_diary_export_item(child)
        if item is not None:
            export_items.append(item)
    if not export_items:
        raise ValueError("没有可导出的日记")

    exporter = DiaryExporter(output_dir)
    docx_path, pdf_path = exporter.export_entries_word_and_pdf(
        export_items,
        export_all=True,
    )
    return {
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
        "output_dir": str(output_dir),
        "count": len(export_items),
    }


def safe_export_name(raw: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*]+', "_", raw)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" ._")
    return cleaned[:100] or "diary_export"


def unique_output_path(target: Path) -> Path:
    if not target.exists():
        return target
    counter = 1
    while True:
        candidate = target.with_name(f"{target.stem}_{counter}{target.suffix}")
        if not candidate.exists():
            return candidate
        counter += 1


def unique_output_dir(target: Path) -> Path:
    if not target.exists():
        return target
    counter = 1
    while True:
        candidate = target.with_name(f"{target.name}_{counter}")
        if not candidate.exists():
            return candidate
        counter += 1


def format_extra_value(value: Any, indent: str = "") -> list[str]:
    lines: list[str] = []
    if isinstance(value, dict):
        for key, item in value.items():
            if item in (None, "", [], {}):
                continue
            if isinstance(item, (dict, list)):
                lines.append(f"{indent}{key}:")
                lines.extend(format_extra_value(item, indent + "  "))
            else:
                lines.append(f"{indent}{key}: {item}")
    elif isinstance(value, list):
        for index, item in enumerate(value, start=1):
            if isinstance(item, (dict, list)):
                lines.append(f"{indent}{index}.")
                lines.extend(format_extra_value(item, indent + "  "))
            elif item not in (None, ""):
                lines.append(f"{indent}{index}. {item}")
    elif value not in (None, ""):
        lines.append(f"{indent}{value}")
    return lines


def module_record_export_lines(module: ModuleConfig, record: dict[str, Any], index: int) -> list[str]:
    body = str(record.get("body") or "").strip()
    lines = [
        f"{index}. {record.get('title') or '未命名记录'}",
        f"日期：{record.get('date') or '未设置'}",
        f"类型：{record.get('type') or record.get('subtitle') or module.label}",
        f"状态：{record.get('status') or '未设置'}",
        f"更新时间：{record.get('updated_at') or '未设置'}",
        f"字数：{len(body)}",
        "",
        body or "（正文为空）",
    ]

    extra = record.get("extra") if isinstance(record.get("extra"), dict) else {}
    detail_lines: list[str] = []
    if module.key == "footprints":
        visits = extra.get("visits") if isinstance(extra.get("visits"), list) else []
        for visit in visits:
            if not isinstance(visit, dict):
                continue
            detail_lines.append(f"- {visit.get('date') or '未设置日期'}")
            thought = str(visit.get("thought") or "").strip()
            if thought:
                detail_lines.append(f"  感想：{thought}")
            images = visit.get("images") if isinstance(visit.get("images"), list) else []
            if images:
                names = []
                for image in images:
                    parsed = image_metadata(image)
                    if parsed:
                        file_name, label = parsed
                        names.append(label or file_name)
                if names:
                    detail_lines.append(f"  图片：{'、'.join(names)}")
    elif module.key == "action_plans":
        tasks = extra.get("tasks") if isinstance(extra.get("tasks"), list) else []
        for task in tasks:
            if not isinstance(task, dict):
                continue
            mark = "完成" if task.get("done") else "未完成"
            title = str(task.get("title") or "未命名子任务")
            task_date = str(task.get("date") or "").strip()
            task_time = str(task.get("time") or "").strip()
            note = str(task.get("note") or "").strip()
            detail_lines.append(f"- [{mark}] {title} {task_date} {task_time}".rstrip())
            if note:
                detail_lines.append(f"  备注：{note}")
    elif module.key == "resources":
        resource_items = extra.get("resource_items") if isinstance(extra.get("resource_items"), list) else []
        for item in resource_items:
            if not isinstance(item, dict):
                continue
            label = str(item.get("type") or item.get("label") or "").strip()
            value = str(item.get("value") or item.get("description") or "").strip()
            if label or value:
                detail_lines.append(f"- {label or '资源'}：{value}")
        judgement = str(extra.get("overall_judgement") or "").strip()
        if judgement:
            detail_lines.append(f"- 综合判断：{judgement}")
    elif module.key == "info_memos":
        fields = extra.get("type_fields") if isinstance(extra.get("type_fields"), dict) else {}
        detail_lines.extend(format_extra_value(fields))
    else:
        images = extra.get("images") if isinstance(extra.get("images"), list) else []
        if images:
            names = []
            for image in images:
                parsed = image_metadata(image)
                if parsed:
                    file_name, label = parsed
                    names.append(label or file_name)
            if names:
                detail_lines.append(f"图片：{'、'.join(names)}")

    if detail_lines:
        lines.extend(["", "补充信息：", *detail_lines])
    lines.extend(["", "-" * 48, ""])
    return lines


def export_module_txt(module_key: str, payload: dict[str, Any] | None = None) -> dict[str, Any]:
    if module_key not in MODULE_BY_KEY:
        raise ValueError(f"未知导出模块：{module_key}")
    module = MODULE_BY_KEY[module_key]
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)
    records = list_module_records(module_key)
    if not records:
        raise ValueError(f"没有可导出的{module.label}记录")

    txt_path = unique_output_path(output_dir / f"{safe_export_name(f'{module.label}_{len(records)}条记录')}.txt")
    lines = [f"{module.label}导出", f"导出时间：{now_iso()}", f"记录数：{len(records)}", ""]
    for index, record in enumerate(records, start=1):
        lines.extend(module_record_export_lines(module, record, index))
    txt_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return {
        "module_key": module.key,
        "module_label": module.label,
        "txt_path": str(txt_path),
        "output_dir": str(output_dir),
        "count": len(records),
    }


def set_docx_runs_font(paragraph: Any, size: int = 11, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(size)
        run.bold = bold


def add_docx_paragraph(document: Any, text: str = "", size: int = 11, bold: bool = False) -> Any:
    paragraph = document.add_paragraph(text)
    set_docx_runs_font(paragraph, size=size, bold=bold)
    return paragraph


def footprint_image_file(place_id: str, visit_id: str, image_name: str) -> Path | None:
    try:
        return footprint_visit_image_path(place_id, visit_id, image_name)
    except Exception:
        return None


def append_footprint_record_to_docx(document: Any, record: dict[str, Any], include_page_break: bool = False) -> None:
    if include_page_break:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    place_id = str(record.get("id") or "")
    title = str(record.get("title") or "未命名足迹")
    heading = document.add_heading(title, level=0)
    set_docx_runs_font(heading, size=18, bold=True)

    body = str(record.get("body") or "").strip()
    add_docx_paragraph(document, f"地点记录数：{len(record.get('extra', {}).get('visits', []) if isinstance(record.get('extra'), dict) else [])}", size=10)
    if record.get("date"):
        add_docx_paragraph(document, f"日期：{record.get('date')}", size=10)
    if body:
        add_docx_paragraph(document, "地点描述", size=13, bold=True)
        for line in body.splitlines():
            add_docx_paragraph(document, line, size=11)

    extra = record.get("extra") if isinstance(record.get("extra"), dict) else {}
    visits = extra.get("visits") if isinstance(extra.get("visits"), list) else []
    if not visits:
        add_docx_paragraph(document, "（暂无访问记录）", size=11)
        return

    for visit in visits:
        if not isinstance(visit, dict):
            continue
        visit_id = str(visit.get("id") or "")
        visit_date = str(visit.get("date") or visit.get("visit_date") or "未设置日期")
        date_heading = document.add_heading(visit_date, level=1)
        set_docx_runs_font(date_heading, size=14, bold=True)

        thought = str(visit.get("thought") or visit.get("reflection") or "").strip()
        if thought:
            for line in thought.splitlines():
                add_docx_paragraph(document, line, size=11)
        else:
            add_docx_paragraph(document, "（本次访问暂无文字记录）", size=11)

        images = visit.get("images") if isinstance(visit.get("images"), list) else []
        exportable_images: list[tuple[str, Path]] = []
        for image in images:
            parsed = image_metadata(image)
            if not parsed:
                continue
            file_name, label = parsed
            image_path = footprint_image_file(place_id, visit_id, file_name)
            if image_path and image_path.exists():
                exportable_images.append((label or file_name, image_path))

        if exportable_images:
            add_docx_paragraph(document, "图片", size=12, bold=True)
            for caption_text, image_path in exportable_images:
                table = document.add_table(rows=1, cols=1)
                cell = table.cell(0, 0)
                caption = cell.paragraphs[0]
                caption.add_run(caption_text)
                set_docx_runs_font(caption, size=10, bold=True)
                image_paragraph = cell.add_paragraph()
                try:
                    image_paragraph.add_run().add_picture(str(image_path), width=Cm(14.8))
                except Exception:
                    image_paragraph.add_run("（图片无法加载）")
                document.add_paragraph("")


def export_footprints_word(payload: dict[str, Any] | None = None, output_dir: Path | None = None) -> dict[str, Any]:
    target_dir = output_dir
    if target_dir is None:
        output_text = str((payload or {}).get("output_dir") or "").strip()
        target_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    target_dir.mkdir(parents=True, exist_ok=True)

    records = list_module_records("footprints")
    if not records:
        raise ValueError("没有可导出的足迹记录")

    docx_path = unique_output_path(target_dir / f"{safe_export_name(f'足迹_{len(records)}条记录')}.docx")
    document = Document()
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"].font.size = Pt(11)
    title = document.add_heading("足迹导出", level=0)
    set_docx_runs_font(title, size=20, bold=True)
    add_docx_paragraph(document, f"导出时间：{now_iso()}", size=10)
    add_docx_paragraph(document, f"地点数：{len(records)}", size=10)

    for index, record in enumerate(records, start=1):
        append_footprint_record_to_docx(document, record, include_page_break=index > 1)

    document.save(docx_path)
    return {
        "module_key": "footprints",
        "module_label": MODULE_BY_KEY["footprints"].label,
        "docx_path": str(docx_path),
        "output_dir": str(target_dir),
        "count": len(records),
    }


def zip_raw_data(target_dir: Path) -> Path:
    zip_path = unique_output_path(target_dir / "2.0原始数据.zip")
    exports_root = (DATA_ROOT / "exports").resolve()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in DATA_ROOT.rglob("*"):
            if not path.is_file():
                continue
            resolved = path.resolve()
            if str(resolved).startswith(str(exports_root)):
                continue
            archive.write(path, path.relative_to(DATA_ROOT))
    return zip_path


def export_all_modules(payload: dict[str, Any] | None = None) -> dict[str, Any]:
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)
    package_dir = unique_output_dir(output_dir / f"全部板块导出_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
    package_dir.mkdir(parents=True, exist_ok=True)

    files: list[dict[str, Any]] = []
    total = 0
    for module in MODULES:
        if module.key == "notes":
            continue
        records = list_module_records(module.key)
        txt_path = package_dir / f"{safe_export_name(module.label)}.txt"
        lines = [f"{module.label}导出", f"导出时间：{now_iso()}", f"记录数：{len(records)}", ""]
        if records:
            for index, record in enumerate(records, start=1):
                lines.extend(module_record_export_lines(module, record, index))
        else:
            lines.append("暂无数据")
        txt_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        total += len(records)
        files.append({"module_key": module.key, "module_label": module.label, "txt_path": str(txt_path), "count": len(records)})
        if module.key == "footprints" and records:
            footprint_word = export_footprints_word(output_dir=package_dir)
            files.append(
                {
                    "module_key": module.key,
                    "module_label": module.label,
                    "docx_path": footprint_word["docx_path"],
                    "count": len(records),
                }
            )

    zip_path = zip_raw_data(package_dir)
    manifest_path = package_dir / "导出说明.json"
    manifest_path.write_text(
        json.dumps(
            {
                "exported_at": now_iso(),
                "data_root": str(DATA_ROOT),
                "total_records": total,
                "files": files,
                "raw_data_zip": str(zip_path),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return {
        "output_dir": str(package_dir),
        "count": total,
        "files": files,
        "zip_path": str(zip_path),
        "manifest_path": str(manifest_path),
    }


def export_all_entries_txt(payload: dict[str, Any] | None = None) -> dict[str, Any]:
    output_text = str((payload or {}).get("output_dir") or "").strip()
    output_dir = Path(output_text).expanduser().resolve() if output_text else configured_export_dir()
    output_dir.mkdir(parents=True, exist_ok=True)

    export_items: list[DiaryExportItem] = []
    entries_dir = DATA_ROOT / "entries"
    entries_dir.mkdir(parents=True, exist_ok=True)
    for child in entries_dir.iterdir():
        if not child.is_dir():
            continue
        item = build_diary_export_item(child)
        if item is not None:
            export_items.append(item)
    if not export_items:
        raise ValueError("没有可导出的日记")

    export_items.sort(
        key=lambda item: (item.entry.date, item.entry.created_at, item.entry.updated_at),
        reverse=True,
    )
    txt_path = unique_output_path(output_dir / f"{safe_export_name(f'全部日记_{len(export_items)}篇日记')}.txt")
    lines: list[str] = []
    for index, item in enumerate(export_items, start=1):
        entry = item.entry
        body = entry.body.strip()
        lines.extend(
            [
                f"{index}. {entry.display_title}",
                f"日期：{entry.date}",
                f"字数：{len(body)}",
                "",
                body or "（正文为空）",
                "",
                "-" * 48,
                "",
            ]
        )
    txt_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return {
        "txt_path": str(txt_path),
        "output_dir": str(output_dir),
        "count": len(export_items),
    }


def update_entry_images(entry_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    record_dir = DATA_ROOT / "entries" / entry_id
    metadata_path = record_dir / "entry.json"
    if not metadata_path.exists():
        raise FileNotFoundError(f"entry not found: {entry_id}")
    images_dir = record_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    requested = payload.get("images")
    if not isinstance(requested, list):
        raise ValueError("images must be a list")

    next_images: list[dict[str, str]] = []
    keep_names: set[str] = set()
    for item in requested:
        parsed = image_metadata(item)
        if not parsed:
            continue
        file_name, label = parsed
        image_path = entry_image_path(entry_id, file_name)
        if not image_path.exists():
            continue
        if file_name in keep_names:
            continue
        next_images.append({"file_name": file_name, "label": label.strip()})
        keep_names.add(file_name)

    data = read_json(metadata_path)
    data["images"] = next_images
    data["updated_at"] = now_iso()
    write_json(metadata_path, data)

    for child in images_dir.iterdir():
        if child.is_file() and child.name not in keep_names:
            child.unlink(missing_ok=True)

    return record_from_directory(MODULE_BY_KEY["entries"], record_dir) or {}


def normalize_images(values: Any, url_builder: Any | None = None) -> list[dict[str, str]]:
    images: list[dict[str, str]] = []
    if not isinstance(values, list):
        return images
    for value in values:
        parsed = image_metadata(value)
        if not parsed:
            continue
        file_name, label = parsed
        image = {"file_name": file_name, "label": label}
        if url_builder is not None:
            image["url"] = str(url_builder(file_name))
        images.append(image)
    return images


def read_footprint_visits(record_dir: Path, data: dict[str, Any]) -> list[dict[str, Any]]:
    visits_dir = record_dir / "visits"
    visits_dir.mkdir(parents=True, exist_ok=True)
    visit_ids = data.get("visit_ids") if isinstance(data.get("visit_ids"), list) else []
    candidates = [visits_dir / str(visit_id) for visit_id in visit_ids]
    candidates.extend(child for child in visits_dir.iterdir() if child.is_dir() and child not in candidates)
    visits: list[dict[str, Any]] = []
    for visit_dir in candidates:
        visit_path = visit_dir / "visit.json"
        if not visit_path.exists():
            continue
        try:
            visit = read_json(visit_path)
        except (OSError, json.JSONDecodeError):
            continue
        if visit.get("deleted"):
            continue
        visit_id = str(visit.get("id") or visit_dir.name)
        thought_path = visit_dir / "thought.md"
        thought = thought_path.read_text(encoding="utf-8", errors="replace") if thought_path.exists() else ""
        visit["id"] = visit_id
        visit["date"] = str(visit.get("date") or visit.get("visit_date") or "")[:10]
        visit["thought"] = thought or str(visit.get("thought") or visit.get("reflection") or "")
        visit["images"] = normalize_images(
            visit.get("images", []),
            lambda name, place_id=record_dir.name, current_visit_id=visit_id: (
                f"/api/modules/footprints/{place_id}/visits/{current_visit_id}/images/{name}"
            ),
        )
        visits.append(visit)
    visits.sort(key=lambda item: str(item.get("date") or item.get("updated_at") or ""), reverse=True)
    return visits


def footprint_visit_dir(place_id: str, visit_id: str) -> Path:
    base = (DATA_ROOT / "footprints" / place_id / "visits").resolve()
    target = (base / visit_id).resolve()
    if not str(target).startswith(str(base)):
        raise ValueError("invalid visit path")
    target.mkdir(parents=True, exist_ok=True)
    return target


def ensure_footprint_visit(place_id: str, visit_date: str) -> tuple[Path, dict[str, Any]]:
    place_dir = DATA_ROOT / "footprints" / place_id
    metadata_path = place_dir / "footprint.json"
    if not metadata_path.exists():
        raise FileNotFoundError(f"footprint not found: {place_id}")
    place_data = read_json(metadata_path)
    visits = read_footprint_visits(place_dir, place_data)
    for visit in visits:
        if str(visit.get("date") or "") == visit_date:
            visit_dir = footprint_visit_dir(place_id, str(visit["id"]))
            return visit_dir, read_json(visit_dir / "visit.json")
    visit_id = uuid.uuid4().hex
    visit_dir = footprint_visit_dir(place_id, visit_id)
    timestamp = now_iso()
    visit = {
        "id": visit_id,
        "date": visit_date,
        "images": [],
        "created_at": timestamp,
        "updated_at": timestamp,
    }
    write_json(visit_dir / "visit.json", visit)
    visit_ids = list(place_data.get("visit_ids", [])) if isinstance(place_data.get("visit_ids"), list) else []
    if visit_id not in visit_ids:
        visit_ids.append(visit_id)
    place_data["visit_ids"] = visit_ids
    place_data["updated_at"] = timestamp
    write_json(metadata_path, place_data)
    return visit_dir, visit


def save_footprint_visit(place_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    visit_date = str(payload.get("date") or date.today().isoformat())[:10]
    visit_dir, visit = ensure_footprint_visit(place_id, visit_date)
    thought = str(payload.get("thought") or "")
    if thought:
        (visit_dir / "thought.md").write_text(thought, encoding="utf-8")
    visit["date"] = visit_date
    visit["updated_at"] = now_iso()
    write_json(visit_dir / "visit.json", visit)
    return record_from_directory(MODULE_BY_KEY["footprints"], DATA_ROOT / "footprints" / place_id) or {}


def footprint_visit_image_path(place_id: str, visit_id: str, image_name: str) -> Path:
    images_root = (DATA_ROOT / "footprints" / place_id / "visits" / visit_id / "images").resolve()
    candidate = (images_root / image_name).resolve()
    if not str(candidate).startswith(str(images_root)):
        raise ValueError("invalid image path")
    if not candidate.exists():
        raise FileNotFoundError(f"image not found: {image_name}")
    return candidate


def add_footprint_visit_images(place_id: str, visit_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    visit_dir = footprint_visit_dir(place_id, visit_id)
    visit_path = visit_dir / "visit.json"
    if not visit_path.exists():
        raise FileNotFoundError(f"visit not found: {visit_id}")
    images_dir = visit_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    visit = read_json(visit_path)
    images = list(visit.get("images", [])) if isinstance(visit.get("images"), list) else []
    files = payload.get("files")
    if not isinstance(files, list):
        raise ValueError("files must be a list")
    for item in files:
        if not isinstance(item, dict):
            continue
        raw_data = str(item.get("data") or "")
        if "," in raw_data and raw_data.startswith("data:"):
            raw_data = raw_data.split(",", 1)[1]
        file_name = unique_image_name(images_dir, str(item.get("name") or "image.png"))
        (images_dir / file_name).write_bytes(base64.b64decode(raw_data))
        images.append({"file_name": file_name, "label": str(item.get("label") or "")})
    visit["images"] = images
    visit["updated_at"] = now_iso()
    write_json(visit_path, visit)
    return record_from_directory(MODULE_BY_KEY["footprints"], DATA_ROOT / "footprints" / place_id) or {}


def update_footprint_visit_images(place_id: str, visit_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    visit_dir = footprint_visit_dir(place_id, visit_id)
    visit_path = visit_dir / "visit.json"
    if not visit_path.exists():
        raise FileNotFoundError(f"visit not found: {visit_id}")
    images_dir = visit_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    requested = payload.get("images")
    if not isinstance(requested, list):
        raise ValueError("images must be a list")
    next_images: list[dict[str, str]] = []
    keep_names: set[str] = set()
    for item in requested:
        parsed = image_metadata(item)
        if not parsed:
            continue
        file_name, label = parsed
        path = footprint_visit_image_path(place_id, visit_id, file_name)
        if path.exists() and file_name not in keep_names:
            next_images.append({"file_name": file_name, "label": label.strip()})
            keep_names.add(file_name)
    visit = read_json(visit_path)
    visit["images"] = next_images
    visit["updated_at"] = now_iso()
    write_json(visit_path, visit)
    for child in images_dir.iterdir():
        if child.is_file() and child.name not in keep_names:
            child.unlink(missing_ok=True)
    return record_from_directory(MODULE_BY_KEY["footprints"], DATA_ROOT / "footprints" / place_id) or {}


def classify_entry_images_to_footprint(entry_id: str, payload: dict[str, Any]) -> dict[str, Any]:
    place_id = str(payload.get("footprint_id") or "")
    visit_date = str(payload.get("date") or date.today().isoformat())[:10]
    selected_names = payload.get("images")
    if not place_id:
        raise ValueError("footprint_id is required")
    if not isinstance(selected_names, list) or not selected_names:
        raise ValueError("images must be a non-empty list")
    entry_dir = DATA_ROOT / "entries" / entry_id
    entry_data = read_json(entry_dir / "entry.json")
    entry_images = normalize_images(entry_data.get("images", []))
    label_by_name = {image["file_name"]: image.get("label", "") for image in entry_images}
    visit_dir, visit = ensure_footprint_visit(place_id, visit_date)
    images_dir = visit_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    images = list(visit.get("images", [])) if isinstance(visit.get("images"), list) else []
    copied = 0
    for image_name in selected_names:
        source = entry_image_path(entry_id, str(image_name))
        file_name = unique_image_name(images_dir, source.name)
        shutil.copy2(source, images_dir / file_name)
        images.append({"file_name": file_name, "label": label_by_name.get(source.name, "")})
        copied += 1
    visit["images"] = images
    visit["updated_at"] = now_iso()
    write_json(visit_dir / "visit.json", visit)
    return {"ok": True, "copied": copied, "footprint": record_from_directory(MODULE_BY_KEY["footprints"], DATA_ROOT / "footprints" / place_id)}


def promote_light_plan_to_action(plan_id: str) -> dict[str, Any]:
    source_dir = DATA_ROOT / "plans" / plan_id
    source_path = source_dir / "plan.json"
    if not source_path.exists():
        raise FileNotFoundError(f"plan not found: {plan_id}")
    source = read_json(source_path)
    body = read_body(source_dir, source, MODULE_BY_KEY["plans"])
    record_id = uuid.uuid4().hex
    record_dir = DATA_ROOT / "action_plans" / record_id
    timestamp = now_iso()
    plan_date = str(source.get("due_date") or source.get("start_date") or date.today().isoformat())[:10]
    data = {
        "id": record_id,
        "title": str(source.get("title") or "未命名行动计划"),
        "plan_type": "日程型行动计划",
        "description": body,
        "start_date": plan_date,
        "end_date": "",
        "status": "进行中",
        "source_light_plan_id": plan_id,
        "tasks": [],
        "summary": "",
        "created_at": timestamp,
        "updated_at": timestamp,
    }
    record_dir.mkdir(parents=True, exist_ok=True)
    write_json(record_dir / "action_plan.json", data)
    return record_from_directory(MODULE_BY_KEY["action_plans"], record_dir) or {}


def save_resource(payload: dict[str, Any]) -> dict[str, Any]:
    record_id = str(payload.get("id") or uuid.uuid4().hex)
    record_dir = DATA_ROOT / "resources" / record_id
    record_dir.mkdir(parents=True, exist_ok=True)
    metadata_path = record_dir / "resource.json"
    existing = read_json(metadata_path) if metadata_path.exists() else {}
    extra = payload.get("extra") if isinstance(payload.get("extra"), dict) else {}
    timestamp = now_iso()
    data = {
        **existing,
        **extra,
        "id": record_id,
        "title": str(payload.get("title") or ""),
        "description": str(payload.get("body") or ""),
        "type": str(extra.get("type") or payload.get("type") or existing.get("type") or "其他"),
        "status": str(extra.get("status") or payload.get("status") or existing.get("status") or "考虑中"),
        "resource_items": extra.get("resource_items") if isinstance(extra.get("resource_items"), list) else existing.get("resource_items", []),
        "overall_judgement": str(extra.get("overall_judgement") or existing.get("overall_judgement") or ""),
        "subjective_feeling": str(extra.get("subjective_feeling") or existing.get("subjective_feeling") or ""),
        "recurrence_test": extra.get("recurrence_test") if isinstance(extra.get("recurrence_test"), dict) else existing.get("recurrence_test", {}),
        "notes": str(extra.get("notes") or existing.get("notes") or ""),
        "related_diaries": existing.get("related_diaries", []),
        "related_thoughts": existing.get("related_thoughts", []),
        "related_self_analysis": existing.get("related_self_analysis", []),
        "created_at": existing.get("created_at") or timestamp,
        "updated_at": timestamp,
    }
    metadata_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return record_from_directory(MODULE_BY_KEY["resources"], record_dir) or {}


def save_generic_record(module_key: str, payload: dict[str, Any]) -> dict[str, Any]:
    if module_key == "entries":
        return save_entry(payload)
    if module_key == "resources":
        return save_resource(payload)

    module = MODULE_BY_KEY[module_key]
    record_id = str(payload.get("id") or uuid.uuid4().hex)
    record_dir = DATA_ROOT / module.directory / record_id
    record_dir.mkdir(parents=True, exist_ok=True)
    metadata_path = record_dir / module.json_file
    existing = read_json(metadata_path) if metadata_path.exists() else {}
    extra = payload.get("extra") if isinstance(payload.get("extra"), dict) else {}
    timestamp = now_iso()
    data = {
        **existing,
        **extra,
        "id": record_id,
        "created_at": existing.get("created_at") or extra.get("created_at") or timestamp,
        "updated_at": timestamp,
    }

    title_field = module.title_fields[0]
    data[title_field] = str(payload.get("title") or data.get(title_field) or "")

    date_value = str(payload.get("date") or "")
    if date_value:
        date_field = module.date_fields[0]
        data[date_field] = date_value

    if payload.get("type"):
        data[module.type_fields[0]] = str(payload.get("type"))
    if payload.get("status"):
        data[module.status_fields[0]] = str(payload.get("status"))

    write_body(record_dir, data, module, str(payload.get("body") or ""))
    if module.key == "footprints":
        data.setdefault("visit_ids", [])
        data.setdefault("images", [])
    if module.key == "action_plans":
        data.setdefault("tasks", [])
    if module.key == "plans":
        data.setdefault("tags", [])
        data.setdefault("plan_type", data.get("plan_type") or "add")

    write_json(metadata_path, data)
    return record_from_directory(module, record_dir) or {}


def delete_generic_record(module_key: str, record_id: str) -> None:
    module = MODULE_BY_KEY[module_key]
    metadata_path = DATA_ROOT / module.directory / record_id / module.json_file
    if not metadata_path.exists():
        raise FileNotFoundError(f"record not found: {record_id}")
    data = read_json(metadata_path)
    timestamp = now_iso()
    data["deleted"] = True
    data["deleted_at"] = timestamp
    data["updated_at"] = timestamp
    write_json(metadata_path, data)


class LifeDiaryHandler(BaseHTTPRequestHandler):
    server_version = "LifeDiary2/1.0"

    def log_message(self, format: str, *args: object) -> None:
        return

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/api/overview":
            self.send_json(build_overview())
            return
        if parsed.path == "/api/settings":
            self.send_json({**read_settings(), "export_dir": str(configured_export_dir())})
            return
        image_match = re.fullmatch(r"/api/modules/entries/([^/]+)/images/(.+)", parsed.path)
        if image_match:
            try:
                path = entry_image_path(unquote(image_match.group(1)), unquote(image_match.group(2)))
                self.send_file(path)
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.NOT_FOUND, str(exc))
            return
        footprint_image_match = re.fullmatch(r"/api/modules/footprints/([^/]+)/visits/([^/]+)/images/(.+)", parsed.path)
        if footprint_image_match:
            try:
                path = footprint_visit_image_path(
                    unquote(footprint_image_match.group(1)),
                    unquote(footprint_image_match.group(2)),
                    unquote(footprint_image_match.group(3)),
                )
                self.send_file(path)
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.NOT_FOUND, str(exc))
            return
        if parsed.path.startswith("/api/modules/"):
            module_key = unquote(parsed.path.removeprefix("/api/modules/"))
            if module_key not in MODULE_BY_KEY:
                self.send_error(HTTPStatus.NOT_FOUND, "unknown module")
                return
            query = parse_qs(parsed.query).get("q", [""])[0]
            self.send_json(list_module_records(module_key, query))
            return
        self.serve_frontend(parsed.path)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/api/actions/open-data-root":
            os.startfile(DATA_ROOT)  # type: ignore[attr-defined]
            self.send_json({"ok": True})
            return
        if parsed.path == "/api/actions/select-export-dir":
            try:
                self.send_json(select_and_store_export_directory())
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        if parsed.path == "/api/settings":
            try:
                self.send_json(write_settings(self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        image_match = re.fullmatch(r"/api/modules/entries/([^/]+)/images", parsed.path)
        if image_match:
            try:
                self.send_json(add_entry_images(unquote(image_match.group(1)), self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        classify_match = re.fullmatch(r"/api/modules/entries/([^/]+)/classify-images", parsed.path)
        if classify_match:
            try:
                self.send_json(classify_entry_images_to_footprint(unquote(classify_match.group(1)), self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        footprint_image_match = re.fullmatch(r"/api/modules/footprints/([^/]+)/visits/([^/]+)/images", parsed.path)
        if footprint_image_match:
            try:
                self.send_json(
                    add_footprint_visit_images(
                        unquote(footprint_image_match.group(1)),
                        unquote(footprint_image_match.group(2)),
                        self.read_json_body(),
                    )
                )
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        visit_match = re.fullmatch(r"/api/modules/footprints/([^/]+)/visits", parsed.path)
        if visit_match:
            try:
                self.send_json(save_footprint_visit(unquote(visit_match.group(1)), self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        promote_match = re.fullmatch(r"/api/modules/plans/([^/]+)/promote-action", parsed.path)
        if promote_match:
            try:
                self.send_json(promote_light_plan_to_action(unquote(promote_match.group(1))))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        if parsed.path == "/api/modules/entries/export-all":
            try:
                self.send_json(export_all_entries_word_pdf(self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        if parsed.path == "/api/modules/entries/export-txt":
            try:
                self.send_json(export_all_entries_txt(self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        if parsed.path == "/api/export/all":
            try:
                self.send_json(export_all_modules(self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        if parsed.path == "/api/modules/footprints/export-word":
            try:
                self.send_json(export_footprints_word(self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        module_export_match = re.fullmatch(r"/api/modules/([^/]+)/export-txt", parsed.path)
        if module_export_match:
            try:
                self.send_json(export_module_txt(unquote(module_export_match.group(1)), self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        export_match = re.fullmatch(r"/api/modules/entries/([^/]+)/export", parsed.path)
        if export_match:
            try:
                self.send_json(export_entry_word_pdf(unquote(export_match.group(1))))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        if not parsed.path.startswith("/api/modules/"):
            self.send_error(HTTPStatus.NOT_FOUND, "unknown endpoint")
            return
        module_key = unquote(parsed.path.removeprefix("/api/modules/"))
        try:
            if module_key not in MODULE_BY_KEY:
                self.send_error(HTTPStatus.NOT_FOUND, "unknown module")
                return
            payload = self.read_json_body()
            self.send_json(save_generic_record(module_key, payload))
        except Exception as exc:  # noqa: BLE001 - this is the app boundary
            self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))

    def do_PUT(self) -> None:
        parsed = urlparse(self.path)
        image_match = re.fullmatch(r"/api/modules/entries/([^/]+)/images", parsed.path)
        if image_match:
            try:
                self.send_json(update_entry_images(unquote(image_match.group(1)), self.read_json_body()))
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        footprint_image_match = re.fullmatch(r"/api/modules/footprints/([^/]+)/visits/([^/]+)/images", parsed.path)
        if footprint_image_match:
            try:
                self.send_json(
                    update_footprint_visit_images(
                        unquote(footprint_image_match.group(1)),
                        unquote(footprint_image_match.group(2)),
                        self.read_json_body(),
                    )
                )
            except Exception as exc:  # noqa: BLE001 - app boundary
                self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))
            return
        self.send_error(HTTPStatus.NOT_FOUND, "unknown endpoint")

    def do_DELETE(self) -> None:
        parsed = urlparse(self.path)
        if not parsed.path.startswith("/api/modules/"):
            self.send_error(HTTPStatus.NOT_FOUND, "unknown endpoint")
            return
        parts = [unquote(part) for part in parsed.path.removeprefix("/api/modules/").split("/") if part]
        if len(parts) != 2 or parts[0] not in MODULE_BY_KEY:
            self.send_error(HTTPStatus.NOT_FOUND, "unknown record")
            return
        try:
            delete_generic_record(parts[0], parts[1])
            self.send_json({"ok": True})
        except Exception as exc:  # noqa: BLE001 - this is the app boundary
            self.send_error(HTTPStatus.INTERNAL_SERVER_ERROR, str(exc))

    def read_json_body(self) -> dict[str, Any]:
        length = int(self.headers.get("Content-Length", "0"))
        raw = self.rfile.read(length).decode("utf-8")
        value = json.loads(raw) if raw else {}
        if not isinstance(value, dict):
            raise ValueError("request body must be an object")
        return value

    def send_json(self, payload: Any) -> None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_file(self, path: Path) -> None:
        content = path.read_bytes()
        mime = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", mime)
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)

    def serve_frontend(self, request_path: str) -> None:
        if not FRONTEND_DIST.exists():
            self.send_error(HTTPStatus.SERVICE_UNAVAILABLE, "frontend dist not found; run npm run build")
            return
        relative = request_path.strip("/") or "index.html"
        target = (FRONTEND_DIST / relative).resolve()
        if not str(target).startswith(str(FRONTEND_DIST.resolve())) or not target.exists() or target.is_dir():
            target = FRONTEND_DIST / "index.html"
        content = target.read_bytes()
        mime = mimetypes.guess_type(str(target))[0] or "application/octet-stream"
        self.send_response(HTTPStatus.OK)
        self.send_header("Content-Type", mime)
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)


def start_server() -> ThreadingHTTPServer:
    migrate_legacy_data_if_needed()
    server = ThreadingHTTPServer(("127.0.0.1", 0), LifeDiaryHandler)
    thread = threading.Thread(target=server.serve_forever, name="LifeDiary2Server", daemon=True)
    thread.start()
    return server


class MainWindow(QMainWindow):
    def __init__(self, url: str):
        super().__init__()
        self.setWindowTitle(APP_TITLE)
        self.resize(1320, 860)
        view = QWebEngineView(self)
        view.load(QUrl(url))
        self.setCentralWidget(view)


def main() -> int:
    server = start_server()
    url = f"http://127.0.0.1:{server.server_port}/"
    app = QApplication(sys.argv)
    window = MainWindow(url)
    window.show()
    result = app.exec()
    server.shutdown()
    return result


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception:
        fallback = FRONTEND_DIST / "index.html"
        if fallback.exists():
            webbrowser.open(fallback.as_uri())
        raise

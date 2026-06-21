from __future__ import annotations

import os
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import DiaryEntry


ProgressCallback = Callable[[int, str], None]


@dataclass(slots=True)
class DiaryExportItem:
    entry: DiaryEntry
    image_lookup: dict[str, Path]


class DiaryExporter:
    def __init__(self, export_root: Path | str):
        self.export_root = Path(export_root)
        self.export_root.mkdir(parents=True, exist_ok=True)

    def export_word_and_pdf(
        self,
        entry: DiaryEntry,
        image_paths: Iterable[Path | str],
        progress: ProgressCallback | None = None,
    ) -> tuple[Path, Path]:
        export_item = DiaryExportItem(
            entry=entry,
            image_lookup=self._build_image_lookup(image_paths),
        )
        return self.export_entries_word_and_pdf([export_item], progress=progress)

    def export_entries_word_and_pdf(
        self,
        export_items: list[DiaryExportItem],
        progress: ProgressCallback | None = None,
        export_all: bool = False,
    ) -> tuple[Path, Path]:
        if not export_items:
            raise ValueError("没有可导出的日记")

        ordered_items = sorted(
            export_items,
            key=lambda item: (item.entry.date, item.entry.created_at, item.entry.updated_at),
            reverse=True,
        )

        callback = progress or (lambda _value, _message: None)
        callback(5, "准备导出路径")

        base_name = self._build_base_name(ordered_items, export_all=export_all)
        docx_path = self._unique_path(self.export_root / f"{base_name}.docx")
        pdf_path = self._unique_path(self.export_root / f"{base_name}.pdf")

        callback(10, "生成 Word 文档")
        self.export_word_entries(ordered_items, docx_path, progress=callback, group_by_date=export_all)

        callback(65, "按 Word 版式生成 PDF")
        self.export_pdf_from_docx(docx_path, pdf_path, progress=callback)

        callback(100, "导出完成")
        return docx_path, pdf_path

    def export_word_entries(
        self,
        export_items: list[DiaryExportItem],
        output_path: Path,
        progress: ProgressCallback | None = None,
        group_by_date: bool = False,
    ) -> Path:
        callback = progress or (lambda _value, _message: None)
        document = Document()
        self._configure_word_styles(document)
        if group_by_date:
            self._append_cover_page(document, export_items)

        total = len(export_items)
        for index, item in enumerate(export_items, start=1):
            callback(
                10 + int(((index - 1) / max(total, 1)) * 50),
                f"写入 Word 日记 {index}/{total}",
            )
            self._append_word_entry(
                document,
                item,
                include_page_break=index > 1,
            )

        callback(60, "保存 Word 文件")
        document.save(output_path)
        return output_path

    def export_pdf_from_docx(
        self,
        docx_path: Path,
        pdf_path: Path,
        progress: ProgressCallback | None = None,
    ) -> Path:
        callback = progress or (lambda _value, _message: None)
        callback(72, "启动 Word 转换 PDF")

        script = """
$ErrorActionPreference = 'Stop'
$docPath = $env:DIARY_EXPORT_DOCX
$pdfPath = $env:DIARY_EXPORT_PDF
$word = $null
$document = $null
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $document = $word.Documents.Open($docPath, $false, $true)
    $document.ExportAsFixedFormat($pdfPath, 17)
    Write-Output 'pdf-export-ok'
}
catch {
    Write-Error $_
    exit 1
}
finally {
    if ($document -ne $null) {
        try { $document.Close($false) } catch {}
    }
    if ($word -ne $null) {
        try { $word.Quit() } catch {}
    }
}
"""

        env = os.environ.copy()
        env["DIARY_EXPORT_DOCX"] = str(docx_path)
        env["DIARY_EXPORT_PDF"] = str(pdf_path)

        result = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-Command",
                script,
            ],
            check=False,
            capture_output=True,
            text=True,
            timeout=180,
            env=env,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )

        callback(95, "完成 PDF 转换")

        # Word 在关闭 COM 对象时偶发抛出清理噪声，但 PDF 已经成功生成。
        # 这里以最终导出的文件是否存在且非空为准，确保 PDF 与 Word 使用同一份版式结果。
        if pdf_path.exists() and pdf_path.stat().st_size > 0:
            return pdf_path

        if result.returncode != 0 or not pdf_path.exists():
            stderr = (result.stderr or "").strip()
            stdout = (result.stdout or "").strip()
            details = stderr or stdout or "Word 转 PDF 失败，但没有返回详细错误。"
            raise RuntimeError(details)

        return pdf_path

    def _append_word_entry(
        self,
        document: Document,
        export_item: DiaryExportItem,
        include_page_break: bool,
    ) -> None:
        entry = export_item.entry
        image_lookup = export_item.image_lookup

        if include_page_break:
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)

        date_heading = document.add_heading(entry.date or "未设置日期", level=1)
        self._style_heading(date_heading, size=Pt(20), color="1F4E79", before=Pt(4), after=Pt(5))

        title = document.add_heading(entry.display_title, level=2)
        self._style_heading(title, size=Pt(15), color="333333", before=Pt(0), after=Pt(3))

        word_count = len(entry.body.strip())
        meta = document.add_paragraph()
        meta.add_run(f"日期：{entry.date or '未设置'}    字数：{word_count}")
        self._set_run_font(meta.runs, size=Pt(9), color="777777")
        self._set_paragraph_format(meta, after=Pt(7), line_spacing=1.1)

        self._add_soft_divider(document)

        if entry.body.strip():
            for line in entry.body.splitlines():
                paragraph = document.add_paragraph(line)
                if line.strip():
                    self._set_run_font(paragraph.runs, size=Pt(11), color="222222")
                    self._set_body_paragraph_format(paragraph)
                else:
                    self._set_paragraph_format(paragraph, after=Pt(2), line_spacing=1.0)
        else:
            paragraph = document.add_paragraph("（正文为空）")
            self._set_run_font(paragraph.runs, size=Pt(11), color="999999")
            self._set_paragraph_format(paragraph, after=Pt(6), line_spacing=1.3)

        exportable_images = []
        for image in entry.images:
            image_path = image_lookup.get(image.file_name)
            if image_path is None:
                continue
            image_path = Path(image_path)
            if not image_path.exists():
                continue
            exportable_images.append((image, image_path))

        if exportable_images:
            heading = document.add_paragraph("图片记录")
            self._set_run_font(heading.runs, size=Pt(12), bold=True, color="1F4E79")
            self._set_paragraph_format(heading, before=Pt(10), after=Pt(6), line_spacing=1.1)
            self._append_image_grid(document, exportable_images)

    def _append_cover_page(self, document: Document, export_items: list[DiaryExportItem]) -> None:
        dates = [item.entry.date for item in export_items if item.entry.date]
        date_range = ""
        if dates:
            date_range = f"{min(dates)} 至 {max(dates)}"

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("人生档案 Diary")
        self._set_run_font(title.runs, size=Pt(26), bold=True, color="1F4E79")
        self._set_paragraph_format(title, before=Pt(120), after=Pt(12), line_spacing=1.0)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("日记导出")
        self._set_run_font(subtitle.runs, size=Pt(16), color="555555")
        self._set_paragraph_format(subtitle, after=Pt(32), line_spacing=1.0)

        info_lines = [
            f"篇数：{len(export_items)}",
            f"时间范围：{date_range or '未设置'}",
        ]
        for text in info_lines:
            paragraph = document.add_paragraph(text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_run_font(paragraph.runs, size=Pt(11), color="666666")
            self._set_paragraph_format(paragraph, after=Pt(4), line_spacing=1.2)

        note = document.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.add_run("本地优先个人记录")
        self._set_run_font(note.runs, size=Pt(10), color="999999")
        self._set_paragraph_format(note, before=Pt(28), after=Pt(0), line_spacing=1.0)

        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _append_image_grid(self, document: Document, exportable_images) -> None:
        columns = 2
        image_width = Cm(7.1)
        cell_width = Cm(7.55)
        for row_start in range(0, len(exportable_images), columns):
            row_images = exportable_images[row_start:row_start + columns]
            table = document.add_table(rows=1, cols=columns)
            table.autofit = False
            self._remove_table_borders(table)
            self._set_row_cant_split(table.rows[0])
            self._set_table_keep_together(table)

            for cell_index, cell in enumerate(table.rows[0].cells):
                self._set_cell_width(cell, cell_width)
                self._set_cell_margins(cell, top=90, bottom=160, left=80, right=80)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                if cell_index >= len(row_images):
                    continue

                image, image_path = row_images[cell_index]
                caption_text = image.label.strip() or "图片"
                display_caption = f"图 {row_start + cell_index + 1}　{caption_text}"

                image_paragraph = cell.paragraphs[0]
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    image_paragraph.add_run().add_picture(str(image_path), width=image_width)
                except Exception:
                    image_paragraph.add_run("（图片无法加载）")
                    self._set_run_font(image_paragraph.runs, size=Pt(10), color="999999")
                self._set_paragraph_format(image_paragraph, after=Pt(4), line_spacing=1.0)
                self._set_keep_with_next(image_paragraph)
                self._set_keep_lines(image_paragraph)

                caption = cell.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.add_run(display_caption)
                self._set_run_font(caption.runs, size=Pt(10.5), bold=True, color="444444")
                self._set_paragraph_format(caption, after=Pt(2), line_spacing=1.0)
                self._set_keep_lines(caption)

            spacer = document.add_paragraph()
            self._set_paragraph_format(spacer, after=Pt(5), line_spacing=1.0)

    def _build_base_name(self, export_items: list[DiaryExportItem], export_all: bool = False) -> str:
        if export_all:
            raw = f"全部日记_{len(export_items)}篇日记"
        elif len(export_items) == 1:
            raw = f"{export_items[0].entry.date}_{export_items[0].entry.display_title}"
        else:
            start_date = export_items[0].entry.date
            end_date = export_items[-1].entry.date
            raw = f"{start_date}_to_{end_date}_{len(export_items)}篇日记"

        cleaned = re.sub(r'[<>:"/\\\\|?*]+', "_", raw)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" ._")
        if not cleaned:
            cleaned = "diary_export"
        return cleaned[:100]

    def _build_image_lookup(self, image_paths: Iterable[Path | str]) -> dict[str, Path]:
        return {
            Path(path).name: Path(path)
            for path in image_paths
        }

    def _unique_path(self, target: Path) -> Path:
        if not target.exists():
            return target

        counter = 1
        while True:
            candidate = target.with_name(f"{target.stem}_{counter}{target.suffix}")
            if not candidate.exists():
                return candidate
            counter += 1

    def _configure_word_styles(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.35)
        section.right_margin = Cm(2.15)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
        self._configure_header_footer(section)

        normal = document.styles["Normal"]
        normal.font.name = "宋体"
        normal.font.size = Pt(11.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.paragraph_format.line_spacing = 1.45
        normal.paragraph_format.space_after = Pt(5)

        for style_name, size, color in (
            ("Heading 1", Pt(20), "1F4E79"),
            ("Heading 2", Pt(16), "2F6F8F"),
            ("Heading 3", Pt(13), "555555"),
            ("Title", Pt(24), "1F4E79"),
        ):
            if style_name in document.styles:
                style = document.styles[style_name]
                style.font.name = "微软雅黑"
                style.font.size = size
                style.font.bold = True
                style.font.color.rgb = RGBColor.from_string(color)
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    def _configure_header_footer(self, section) -> None:
        header = section.header.paragraphs[0]
        header.text = "人生档案 Diary · 日记导出"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_run_font(header.runs, size=Pt(9), color="888888")
        self._set_paragraph_format(header, after=Pt(0), line_spacing=1.0)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_run_font([footer.add_run("第 ")], size=Pt(9), color="888888")
        page_run = footer.add_run()
        self._add_page_number_field(page_run)
        self._set_run_font([page_run], size=Pt(9), color="888888")
        self._set_run_font([footer.add_run(" 页")], size=Pt(9), color="888888")
        self._set_paragraph_format(footer, before=Pt(0), after=Pt(0), line_spacing=1.0)

    def _set_run_font(self, runs, size: Pt, bold: bool = False, color: str | None = None) -> None:
        for run in runs:
            run.font.name = "宋体"
            run.font.size = size
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _style_heading(self, paragraph, size: Pt, color: str, before: Pt, after: Pt) -> None:
        self._set_run_font(paragraph.runs, size=size, bold=True, color=color)
        for run in paragraph.runs:
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        self._set_paragraph_format(paragraph, before=before, after=after, line_spacing=1.1)

    def _set_paragraph_format(self, paragraph, before: Pt | None = None, after: Pt | None = None, line_spacing: float | None = None) -> None:
        if before is not None:
            paragraph.paragraph_format.space_before = before
        if after is not None:
            paragraph.paragraph_format.space_after = after
        if line_spacing is not None:
            paragraph.paragraph_format.line_spacing = line_spacing

    def _set_body_paragraph_format(self, paragraph) -> None:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.45
        paragraph.paragraph_format.space_after = Pt(5)

    def _add_soft_divider(self, document: Document) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        p_pr = paragraph._element.get_or_add_pPr()
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D9E2EC")

    def _set_keep_with_next(self, paragraph) -> None:
        """Set keepWithNext on a paragraph to keep it on the same page as the next paragraph."""
        pPr = paragraph._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = paragraph._element.makeelement(qn("w:pPr"), {})
            paragraph._element.insert(0, pPr)
        keep = pPr.find(qn("w:keepNext"))
        if keep is None:
            keep = pPr.makeelement(qn("w:keepNext"), {})
            pPr.append(keep)

    def _set_keep_lines(self, paragraph) -> None:
        """Set keepLines on a paragraph to keep all its lines on the same page."""
        pPr = paragraph._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = paragraph._element.makeelement(qn("w:pPr"), {})
            paragraph._element.insert(0, pPr)
        keep = pPr.find(qn("w:keepLines"))
        if keep is None:
            keep = pPr.makeelement(qn("w:keepLines"), {})
            pPr.append(keep)

    def _set_row_cant_split(self, row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    def _set_table_keep_together(self, table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._set_keep_lines(paragraph)

    def _set_cell_width(self, cell, width: Cm) -> None:
        cell.width = width
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(int(width.twips)))
        tc_w.set(qn("w:type"), "dxa")

    def _set_cell_margins(self, cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        margins = tc_pr.find(qn("w:tcMar"))
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tc_pr.append(margins)
        for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
            element = margins.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                margins.append(element)
            element.set(qn("w:w"), str(value))
            element.set(qn("w:type"), "dxa")

    def _add_page_number_field(self, run) -> None:
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    def _remove_table_borders(self, table) -> None:
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "nil")
